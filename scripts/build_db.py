"""
GTİP Veritabanı Oluşturucu
===========================
Türk Gümrük Tarife Cetveli fasıl XLS dosyalarını, fasıl notlarını,
izahname notlarını, yorum kurallarını ve içindekiler haritasını
parse edip SQLite veritabanına yazar.

Kullanım:
    # Temel (fasıl + notlar)
    python build_db.py data/fasil_dosyalari/ --notlar data/fasil_notlari/

    # Tam build (Faz 2)
    python build_db.py data/fasil_dosyalari/ \\
        --notlar data/fasil_notlari/ \\
        --izahname data/izahname_notlari/ \\
        --yorum data/yorum_kurallari/ \\
        --icindekiler data/icindekiler/ \\
        --db data/gtip_2026.db --force
"""

import sys
import os
import re
import sqlite3
import glob
import argparse
import subprocess
import tempfile

try:
    import xlrd
except ImportError:
    print("xlrd yüklü değil: pip install xlrd")
    sys.exit(1)


# ---------------------------------------------------------------------------
# DB Schema
# ---------------------------------------------------------------------------

def create_db(db_path):
    """SQLite veritabanı şemasını oluştur (mevcut tabloları siler)."""
    conn = sqlite3.connect(db_path)
    c = conn.cursor()

    c.execute("DROP TABLE IF EXISTS gtip_fts")
    c.execute("DROP TABLE IF EXISTS notlar_fts")
    c.execute("DROP TABLE IF EXISTS izahname_fts")
    c.execute("DROP TABLE IF EXISTS gtip")
    c.execute("DROP TABLE IF EXISTS pozisyon")
    c.execute("DROP TABLE IF EXISTS fasil_meta")
    c.execute("DROP TABLE IF EXISTS fasil_notlari")
    c.execute("DROP TABLE IF EXISTS izahname_notlari")
    c.execute("DROP TABLE IF EXISTS yorum_kurallari")
    c.execute("DROP TABLE IF EXISTS bolum_fasil")

    c.execute("""
        CREATE TABLE gtip (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            gtip_code TEXT NOT NULL,
            gtip_clean TEXT NOT NULL,
            fasil INTEGER,
            pozisyon TEXT,
            alt_pozisyon TEXT,
            tanim TEXT NOT NULL,
            tanim_hiyerarsi TEXT,
            olcu_birimi TEXT,
            seviye INTEGER,
            UNIQUE(gtip_code)
        )
    """)

    c.execute("""
        CREATE TABLE pozisyon (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            kod TEXT NOT NULL,
            kod_clean TEXT NOT NULL,
            fasil INTEGER,
            tanim TEXT NOT NULL,
            seviye INTEGER,
            UNIQUE(kod)
        )
    """)

    c.execute("""
        CREATE TABLE fasil_meta (
            fasil_no INTEGER PRIMARY KEY,
            dosya_adi TEXT,
            satir_sayisi INTEGER,
            gtip_sayisi INTEGER,
            parse_tarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

    c.execute("""
        CREATE TABLE fasil_notlari (
            fasil_no INTEGER PRIMARY KEY,
            bolum_notu TEXT,
            fasil_notu TEXT,
            tam_metin TEXT,
            kelime_sayisi INTEGER,
            dosya_adi TEXT
        )
    """)

    c.execute("""
        CREATE TABLE izahname_notlari (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fasil_no INTEGER,
            pozisyon TEXT,
            metin TEXT,
            kelime_sayisi INTEGER,
            dosya_adi TEXT
        )
    """)

    c.execute("""
        CREATE TABLE yorum_kurallari (
            kural_no INTEGER PRIMARY KEY,
            baslik TEXT,
            metin TEXT
        )
    """)

    c.execute("""
        CREATE TABLE bolum_fasil (
            bolum_no INTEGER,
            bolum_adi TEXT,
            fasil_no INTEGER,
            fasil_adi TEXT
        )
    """)

    c.execute("CREATE INDEX idx_gtip_clean ON gtip(gtip_clean)")
    c.execute("CREATE INDEX idx_gtip_fasil ON gtip(fasil)")
    c.execute("CREATE INDEX idx_gtip_pozisyon ON gtip(pozisyon)")
    c.execute("CREATE INDEX idx_gtip_alt_pozisyon ON gtip(alt_pozisyon)")
    c.execute("CREATE INDEX idx_pozisyon_kod ON pozisyon(kod_clean)")
    c.execute("CREATE INDEX idx_izahname_fasil ON izahname_notlari(fasil_no)")

    c.execute("""
        CREATE VIRTUAL TABLE gtip_fts
        USING fts5(gtip_code, tanim, tanim_hiyerarsi, content=gtip, content_rowid=id)
    """)

    c.execute("""
        CREATE VIRTUAL TABLE notlar_fts
        USING fts5(fasil_no, tam_metin)
    """)

    c.execute("""
        CREATE VIRTUAL TABLE izahname_fts
        USING fts5(fasil_no, pozisyon, metin, content=izahname_notlari, content_rowid=id)
    """)

    conn.commit()
    return conn


# ---------------------------------------------------------------------------
# Fasıl (tarife) parser
# ---------------------------------------------------------------------------

def parse_fasil_xls(filepath):
    wb = xlrd.open_workbook(filepath)
    sh = wb.sheet_by_index(0)

    entries = []
    for r in range(sh.nrows):
        pos = str(sh.cell_value(r, 0)).strip()
        desc = str(sh.cell_value(r, 1)).strip()
        olcu = str(sh.cell_value(r, 2)).strip() if sh.ncols > 2 else ""

        if olcu and olcu.replace('.', '').replace('0', '') == '':
            olcu = ""

        entries.append((pos, desc, olcu))

    merged = []
    for pos, desc, olcu in entries:
        if pos == "" and desc != "" and merged:
            # Tire ile başlayan satır → hayalet başlık, bir öncekine merge etme
            if desc.lstrip().startswith('-'):
                merged.append(("", desc, olcu))
            else:
                prev = merged[-1]
                merged[-1] = (prev[0], prev[1] + " " + desc, prev[2] or olcu)
        elif pos != "" or desc != "":
            merged.append((pos, desc, olcu))

    fasil_no = None
    m = re.search(r'(\d+)', os.path.basename(filepath))
    if m:
        fasil_no = int(m.group(1))

    gtip_rows = []
    pozisyon_rows = []
    hierarchy = {}

    for pos, desc, olcu in merged:
        clean = pos.replace(".", "").replace(" ", "")

        # Hayalet başlık: pos boş, desc tire ile başlıyor
        if not clean and desc.lstrip().startswith('-'):
            dash_count = 0
            for ch in desc:
                if ch == '-':
                    dash_count += 1
                elif ch != ' ':
                    break
            clean_desc = re.sub(r'^[\s\-]+', '', desc).strip()
            if clean_desc.endswith(':'):
                clean_desc = clean_desc[:-1].strip()
            hierarchy[dash_count] = clean_desc
            for k in list(hierarchy.keys()):
                if k > dash_count:
                    del hierarchy[k]
            continue

        if not clean.isdigit() or len(clean) < 4:
            continue

        dash_count = 0
        for ch in desc:
            if ch == '-':
                dash_count += 1
            elif ch != ' ':
                break

        clean_desc = re.sub(r'^[\s\-]+', '', desc).strip()
        if clean_desc.endswith(':'):
            clean_desc = clean_desc[:-1].strip()

        if len(clean) == 12:
            pozisyon_4 = clean[:4]
            alt_poz_6 = clean[:6]

            parts = []
            for lvl in sorted(hierarchy.keys()):
                if lvl < dash_count:
                    parts.append(hierarchy[lvl])
            parts.append(clean_desc)
            tanim_hiyerarsi = " > ".join(parts)

            gtip_rows.append({
                'gtip_code': pos,
                'gtip_clean': clean,
                'fasil': fasil_no,
                'pozisyon': pozisyon_4[:2] + "." + pozisyon_4[2:],
                'alt_pozisyon': alt_poz_6,
                'tanim': clean_desc,
                'tanim_hiyerarsi': tanim_hiyerarsi,
                'olcu_birimi': olcu if olcu != '-' else '',
                'seviye': dash_count
            })
        elif len(clean) in (4, 6, 8, 10):
            pozisyon_rows.append({
                'kod': pos,
                'kod_clean': clean,
                'fasil': fasil_no,
                'tanim': clean_desc,
                'seviye': len(clean)
            })

            hierarchy[dash_count] = clean_desc
            for k in list(hierarchy.keys()):
                if k > dash_count:
                    del hierarchy[k]

    return fasil_no, gtip_rows, pozisyon_rows


# ---------------------------------------------------------------------------
# Fasıl notları parser
# ---------------------------------------------------------------------------

def parse_fasil_notu(filepath):
    wb = xlrd.open_workbook(filepath)
    sh = wb.sheet_by_index(0)

    lines = []
    for r in range(sh.nrows):
        row_parts = []
        for c_idx in range(sh.ncols):
            v = str(sh.cell_value(r, c_idx)).strip()
            if v:
                row_parts.append(v)
        if row_parts:
            lines.append(" ".join(row_parts))
        else:
            lines.append("")

    fasil_no = None
    m = re.search(r'(\d+)', os.path.basename(filepath))
    if m:
        fasil_no = int(m.group(1))

    fasil_line_idx = None
    for i, line in enumerate(lines):
        if re.match(r'^FAS[IİıiÝ]L\s+\S', line, re.IGNORECASE):
            fasil_line_idx = i
            break

    if fasil_line_idx is not None and fasil_line_idx > 0:
        bolum_text = "\n".join(lines[:fasil_line_idx]).strip()
        fasil_text = "\n".join(lines[fasil_line_idx:]).strip()
    else:
        bolum_text = ""
        fasil_text = "\n".join(lines).strip()

    tam_metin = (bolum_text + "\n" + fasil_text).strip() if bolum_text else fasil_text

    return fasil_no, bolum_text, fasil_text, tam_metin


# ---------------------------------------------------------------------------
# .doc → .docx dönüştürücü
# ---------------------------------------------------------------------------

def doc_to_docx(doc_path, out_dir):
    """soffice ile .doc → .docx dönüştür. Döner: .docx dosya yolu."""
    result = subprocess.run(
        ['soffice', '--headless', '--convert-to', 'docx',
         os.path.abspath(doc_path), '--outdir', out_dir],
        capture_output=True, text=True
    )
    basename = os.path.splitext(os.path.basename(doc_path))[0]
    docx_path = os.path.join(out_dir, basename + '.docx')
    if not os.path.exists(docx_path):
        raise RuntimeError(f"soffice dönüşümü başarısız: {result.stderr.strip()}")
    return docx_path


# ---------------------------------------------------------------------------
# İzahname notları parser
# ---------------------------------------------------------------------------

def parse_izahname_doc(filepath, tmpdir):
    """
    Tek bir izahname .doc dosyasını parse et.
    Returns: (fasil_no, metin, kelime_sayisi)
    Metin: tüm paragraflar birleştirilmiş tam izahname metni.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx yüklü değil: pip install python-docx")

    fasil_no = None
    m = re.search(r'(\d+)', os.path.basename(filepath))
    if m:
        fasil_no = int(m.group(1))

    docx_path = doc_to_docx(filepath, tmpdir)
    doc = Document(docx_path)

    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    metin = '\n'.join(lines)
    kelime_sayisi = len(metin.split())

    return fasil_no, metin, kelime_sayisi


# ---------------------------------------------------------------------------
# Yorum kuralları parser
# ---------------------------------------------------------------------------

def parse_yorum_kurallari_doc(filepath, tmpdir):
    """
    Yorum kuralları .doc dosyasını parse et.
    Returns: list of (kural_no, baslik, metin)
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx yüklü değil: pip install python-docx")

    docx_path = doc_to_docx(filepath, tmpdir)
    doc = Document(docx_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    kural_re = re.compile(r'^KURAL\s+(\d+)\s*$', re.IGNORECASE)

    rules = []
    current_no = None
    current_baslik = ''
    current_lines = []

    for line in paragraphs:
        m = kural_re.match(line)
        if m:
            if current_no is not None:
                rules.append((current_no, current_baslik, '\n'.join(current_lines)))
            current_no = int(m.group(1))
            current_baslik = line
            current_lines = []
        elif current_no is not None:
            current_lines.append(line)

    if current_no is not None:
        rules.append((current_no, current_baslik, '\n'.join(current_lines)))

    return rules


# ---------------------------------------------------------------------------
# İçindekiler parser
# ---------------------------------------------------------------------------

_ROMAN_VAL = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}

def _roman_to_int(s):
    s = s.upper().strip()
    total, prev = 0, 0
    for ch in reversed(s):
        val = _ROMAN_VAL.get(ch, 0)
        total += val if val >= prev else -val
        prev = val
    return total


def parse_icindekiler_xls(filepath):
    """
    İçindekiler XLS dosyasını parse et.
    Returns: list of (bolum_no, bolum_adi, fasil_no, fasil_adi)
    """
    wb = xlrd.open_workbook(filepath)
    sh = wb.sheet_by_index(0)

    rows = []
    for r in range(sh.nrows):
        row = [str(sh.cell_value(r, c)).strip() for c in range(sh.ncols)]
        row = [x for x in row if x]
        rows.append(row)

    bolum_re = re.compile(r'^BÖLÜM\s+([IVXLCDM]+)\s*$', re.IGNORECASE)

    result = []
    current_bolum_no = None
    current_bolum_adi_parts = []
    in_bolum_adi = False
    current_fasil_no = None
    current_fasil_adi_parts = []

    def flush_fasil():
        if current_fasil_no is not None and current_bolum_no is not None:
            adi = ' '.join(current_fasil_adi_parts)
            adi = re.sub(r'[.…]{2,}.*$', '', adi).strip()
            bolum_adi = ' '.join(current_bolum_adi_parts)
            result.append((current_bolum_no, bolum_adi, current_fasil_no, adi))

    for row in rows:
        if not row:
            continue

        text = ' '.join(row)

        m = bolum_re.match(text)
        if m:
            flush_fasil()
            current_fasil_no = None
            current_fasil_adi_parts = []
            current_bolum_no = _roman_to_int(m.group(1))
            current_bolum_adi_parts = []
            in_bolum_adi = True
            continue

        if in_bolum_adi:
            # "Fasıl No:" satırı bölüm adının bitişini işaret eder
            if re.match(r'^fas[iı]l\s+no', text, re.IGNORECASE):
                in_bolum_adi = False
                continue
            # Rakamla başlıyorsa bölüm adı bitmiş, fasıl başlamış demektir
            if row[0].isdigit():
                in_bolum_adi = False
                # fasıl satırını işle (fall through)
            else:
                current_bolum_adi_parts.append(text)
                continue

        # Fasıl No: başlığını atla
        if re.match(r'^fas[iı]l\s+no', text, re.IGNORECASE):
            continue

        if row[0].isdigit():
            flush_fasil()
            current_fasil_no = int(row[0])
            current_fasil_adi_parts = row[1:] if len(row) > 1 else []
        elif current_fasil_no is not None:
            # Bir önceki fasıl adının devamı
            current_fasil_adi_parts.extend(row)

    flush_fasil()
    return result


# ---------------------------------------------------------------------------
# DB insert
# ---------------------------------------------------------------------------

def insert_tarife(conn, fasil_no, gtip_rows, pozisyon_rows, filename):
    c = conn.cursor()
    c.execute("DELETE FROM gtip WHERE fasil = ?", (fasil_no,))
    c.execute("DELETE FROM pozisyon WHERE fasil = ?", (fasil_no,))
    c.execute("DELETE FROM fasil_meta WHERE fasil_no = ?", (fasil_no,))

    for row in gtip_rows:
        c.execute("""
            INSERT OR REPLACE INTO gtip
            (gtip_code, gtip_clean, fasil, pozisyon, alt_pozisyon,
             tanim, tanim_hiyerarsi, olcu_birimi, seviye)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            row['gtip_code'], row['gtip_clean'], row['fasil'],
            row['pozisyon'], row['alt_pozisyon'], row['tanim'],
            row['tanim_hiyerarsi'], row['olcu_birimi'], row['seviye']
        ))

    for row in pozisyon_rows:
        c.execute("""
            INSERT OR REPLACE INTO pozisyon (kod, kod_clean, fasil, tanim, seviye)
            VALUES (?, ?, ?, ?, ?)
        """, (row['kod'], row['kod_clean'], row['fasil'], row['tanim'], row['seviye']))

    c.execute("""
        INSERT OR REPLACE INTO fasil_meta (fasil_no, dosya_adi, satir_sayisi, gtip_sayisi)
        VALUES (?, ?, ?, ?)
    """, (fasil_no, filename, len(gtip_rows) + len(pozisyon_rows), len(gtip_rows)))

    conn.commit()
    return len(gtip_rows), len(pozisyon_rows)


def insert_notlar(conn, fasil_no, bolum_notu, fasil_notu, tam_metin, filename):
    c = conn.cursor()
    c.execute("DELETE FROM fasil_notlari WHERE fasil_no = ?", (fasil_no,))
    kelime = len(tam_metin.split())
    c.execute("""
        INSERT INTO fasil_notlari (fasil_no, bolum_notu, fasil_notu, tam_metin, kelime_sayisi, dosya_adi)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (fasil_no, bolum_notu, fasil_notu, tam_metin, kelime, filename))
    conn.commit()
    return kelime


def insert_izahname(conn, fasil_no, metin, kelime_sayisi, filename):
    c = conn.cursor()
    c.execute("DELETE FROM izahname_notlari WHERE fasil_no = ? AND pozisyon IS NULL", (fasil_no,))
    c.execute("""
        INSERT INTO izahname_notlari (fasil_no, pozisyon, metin, kelime_sayisi, dosya_adi)
        VALUES (?, NULL, ?, ?, ?)
    """, (fasil_no, metin, kelime_sayisi, filename))
    conn.commit()


def insert_yorum_kurallari(conn, rules):
    c = conn.cursor()
    c.execute("DELETE FROM yorum_kurallari")
    for kural_no, baslik, metin in rules:
        c.execute("""
            INSERT OR REPLACE INTO yorum_kurallari (kural_no, baslik, metin)
            VALUES (?, ?, ?)
        """, (kural_no, baslik, metin))
    conn.commit()


def insert_bolum_fasil(conn, rows):
    c = conn.cursor()
    c.execute("DELETE FROM bolum_fasil")
    c.executemany("""
        INSERT INTO bolum_fasil (bolum_no, bolum_adi, fasil_no, fasil_adi)
        VALUES (?, ?, ?, ?)
    """, rows)
    conn.commit()


def rebuild_fts(conn):
    """FTS indekslerini yeniden oluştur."""
    c = conn.cursor()
    c.execute("INSERT INTO gtip_fts(gtip_fts) VALUES('rebuild')")

    c.execute("DELETE FROM notlar_fts")
    rows = c.execute("SELECT fasil_no, tam_metin FROM fasil_notlari").fetchall()
    for fno, metin in rows:
        c.execute("INSERT INTO notlar_fts (fasil_no, tam_metin) VALUES (?, ?)", (str(fno), metin))

    c.execute("INSERT INTO izahname_fts(izahname_fts) VALUES('rebuild')")

    conn.commit()


# ---------------------------------------------------------------------------
# Stats
# ---------------------------------------------------------------------------

def print_stats(conn):
    c = conn.cursor()

    total_gtip    = c.execute("SELECT COUNT(*) FROM gtip").fetchone()[0]
    total_poz     = c.execute("SELECT COUNT(*) FROM pozisyon").fetchone()[0]
    total_fasil   = c.execute("SELECT COUNT(*) FROM fasil_meta").fetchone()[0]
    total_notlar  = c.execute("SELECT COUNT(*) FROM fasil_notlari").fetchone()[0]
    total_izahname = c.execute("SELECT COUNT(*) FROM izahname_notlari").fetchone()[0]
    total_kural   = c.execute("SELECT COUNT(*) FROM yorum_kurallari").fetchone()[0]
    total_bf      = c.execute("SELECT COUNT(*) FROM bolum_fasil").fetchone()[0]
    total_kelime  = c.execute("SELECT COALESCE(SUM(kelime_sayisi), 0) FROM fasil_notlari").fetchone()[0]
    izah_kelime   = c.execute("SELECT COALESCE(SUM(kelime_sayisi), 0) FROM izahname_notlari").fetchone()[0]

    print(f"\n{'='*50}")
    print(f"VERİTABANI ÖZETİ")
    print(f"{'='*50}")
    print(f"Toplam fasıl       : {total_fasil}")
    print(f"Toplam GTİP        : {total_gtip}")
    print(f"Toplam pozisyon    : {total_poz}")
    print(f"Fasıl notları      : {total_notlar} ({total_kelime:,} kelime)")
    print(f"İzahname notları   : {total_izahname} fasıl ({izah_kelime:,} kelime)")
    print(f"Yorum kuralları    : {total_kural} kural")
    print(f"Bölüm-fasıl harita : {total_bf} satır")
    print(f"{'='*50}")

    print(f"\nFasıl detayları:")
    rows = c.execute("""
        SELECT fasil_no, dosya_adi, gtip_sayisi
        FROM fasil_meta ORDER BY fasil_no
    """).fetchall()
    for fasil_no, dosya, gtip_sayi in rows:
        print(f"  Fasıl {fasil_no:2d}: {gtip_sayi:5d} GTİP  ({dosya})")

    print(f"\nÖrnek arama: 'plastik':")
    rows = c.execute("""
        SELECT gtip_code, tanim FROM gtip_fts
        WHERE gtip_fts MATCH 'plastik' LIMIT 5
    """).fetchall()
    for code, tanim in rows:
        print(f"  {code}: {tanim[:70]}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='GTİP Veritabanı Oluşturucu',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Örnekler:
  python build_db.py data/fasil_dosyalari/ --notlar data/fasil_notlari/
  python build_db.py data/fasil_dosyalari/ \\
      --notlar data/fasil_notlari/ \\
      --izahname data/izahname_notlari/ \\
      --yorum data/yorum_kurallari/ \\
      --icindekiler data/icindekiler/ \\
      --db data/gtip_2026.db --force
        """
    )
    parser.add_argument('input', help='Fasıl XLS dosyası veya klasör yolu')
    parser.add_argument('--notlar',      help='Fasıl notları klasör yolu')
    parser.add_argument('--izahname',    help='İzahname notları klasör yolu (.doc dosyaları)')
    parser.add_argument('--yorum',       help='Yorum kuralları klasör yolu (.doc veya .xls)')
    parser.add_argument('--icindekiler', help='İçindekiler klasör yolu (.xls)')
    parser.add_argument('--db',    default='data/gtip_2026.db', help='SQLite DB çıktı yolu')
    parser.add_argument('--force', action='store_true', help='Mevcut DB varsa üzerine yaz')
    args = parser.parse_args()

    db_path = args.db

    if os.path.exists(db_path) and not args.force:
        print(f"Hata: {db_path} zaten mevcut. Üzerine yazmak için --force kullanın.")
        sys.exit(1)

    # --- Dosya listelerini topla ---
    if os.path.isdir(args.input):
        files = sorted(
            glob.glob(os.path.join(args.input, '*.xls')) +
            glob.glob(os.path.join(args.input, '*.xlsx'))
        )
    elif os.path.isfile(args.input):
        files = [args.input]
    else:
        print(f"Hata: {args.input} bulunamadı")
        sys.exit(1)

    if not files:
        print(f"Hata: {args.input} içinde XLS dosyası bulunamadı")
        sys.exit(1)

    note_files = []
    if args.notlar:
        if os.path.isdir(args.notlar):
            note_files = sorted(
                glob.glob(os.path.join(args.notlar, '*.xls')) +
                glob.glob(os.path.join(args.notlar, '*.xlsx'))
            )
        else:
            print(f"Uyarı: {args.notlar} bulunamadı, notlar atlanıyor.")

    izahname_files = []
    if args.izahname:
        if os.path.isdir(args.izahname):
            izahname_files = sorted(glob.glob(os.path.join(args.izahname, '*.doc')))
        else:
            print(f"Uyarı: {args.izahname} bulunamadı, izahname atlanıyor.")

    yorum_doc = None
    if args.yorum:
        if os.path.isdir(args.yorum):
            docs = glob.glob(os.path.join(args.yorum, '*.doc'))
            if docs:
                yorum_doc = docs[0]
        elif os.path.isfile(args.yorum):
            yorum_doc = args.yorum

    icindekiler_xls = None
    if args.icindekiler:
        if os.path.isdir(args.icindekiler):
            xlss = (glob.glob(os.path.join(args.icindekiler, '*.xls')) +
                    glob.glob(os.path.join(args.icindekiler, '*.xlsx')))
            if xlss:
                icindekiler_xls = xlss[0]
        elif os.path.isfile(args.icindekiler):
            icindekiler_xls = args.icindekiler

    print(f"Veritabanı      : {db_path}")
    print(f"Fasıl dosyası   : {len(files)}")
    print(f"Not dosyası     : {len(note_files)}")
    print(f"İzahname dosyası: {len(izahname_files)}")
    print(f"Yorum kuralları : {os.path.basename(yorum_doc) if yorum_doc else '-'}")
    print(f"İçindekiler     : {os.path.basename(icindekiler_xls) if icindekiler_xls else '-'}")
    print()

    conn = create_db(db_path)
    errors = []

    # --- Tarife parse ---
    total_gtip = 0
    total_poz = 0
    for filepath in files:
        filename = os.path.basename(filepath)
        try:
            fasil_no, gtip_rows, pozisyon_rows = parse_fasil_xls(filepath)
            n_gtip, n_poz = insert_tarife(conn, fasil_no, gtip_rows, pozisyon_rows, filename)
            total_gtip += n_gtip
            total_poz += n_poz
            print(f"  + Fasıl {fasil_no:2d}: {n_gtip:5d} GTİP, {n_poz:3d} pozisyon  ({filename})")
        except Exception as e:
            errors.append((filename, str(e)))
            print(f"  ! HATA: {filename} -> {e}")

    # --- Notlar parse ---
    total_kelime = 0
    for filepath in note_files:
        filename = os.path.basename(filepath)
        try:
            fasil_no, bolum, fasil, tam = parse_fasil_notu(filepath)
            kelime = insert_notlar(conn, fasil_no, bolum, fasil, tam, filename)
            total_kelime += kelime
            print(f"  + Not {fasil_no:2d}: {kelime:5d} kelime  ({filename})")
        except Exception as e:
            errors.append((filename, str(e)))
            print(f"  ! HATA: {filename} -> {e}")

    # --- İzahname parse (.doc → docx dönüşümü için tmpdir) ---
    if izahname_files:
        print(f"\nİzahname parse ediliyor ({len(izahname_files)} dosya)...")
        with tempfile.TemporaryDirectory() as tmpdir:
            for filepath in izahname_files:
                filename = os.path.basename(filepath)
                try:
                    fasil_no, metin, kelime_sayisi = parse_izahname_doc(filepath, tmpdir)
                    insert_izahname(conn, fasil_no, metin, kelime_sayisi, filename)
                    print(f"  + İzahname {fasil_no:2d}: {kelime_sayisi:6d} kelime  ({filename})")
                except Exception as e:
                    errors.append((filename, str(e)))
                    print(f"  ! HATA: {filename} -> {e}")

    # --- Yorum kuralları parse ---
    if yorum_doc:
        print(f"\nYorum kuralları parse ediliyor: {os.path.basename(yorum_doc)}")
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                rules = parse_yorum_kurallari_doc(yorum_doc, tmpdir)
                insert_yorum_kurallari(conn, rules)
                print(f"  + {len(rules)} kural yüklendi")
            except Exception as e:
                errors.append((os.path.basename(yorum_doc), str(e)))
                print(f"  ! HATA: {e}")

    # --- İçindekiler parse ---
    if icindekiler_xls:
        print(f"\nİçindekiler parse ediliyor: {os.path.basename(icindekiler_xls)}")
        try:
            bf_rows = parse_icindekiler_xls(icindekiler_xls)
            insert_bolum_fasil(conn, bf_rows)
            print(f"  + {len(bf_rows)} fasıl-bölüm kaydı yüklendi")
        except Exception as e:
            errors.append((os.path.basename(icindekiler_xls), str(e)))
            print(f"  ! HATA: {e}")

    # --- FTS rebuild ---
    rebuild_fts(conn)

    # --- Sonuç ---
    print_stats(conn)

    if errors:
        print(f"\n! {len(errors)} dosyada hata oluştu:")
        for fname, err in errors:
            print(f"  {fname}: {err}")

    conn.close()
    print(f"\nVeritabanı kaydedildi: {db_path}")


if __name__ == "__main__":
    main()
